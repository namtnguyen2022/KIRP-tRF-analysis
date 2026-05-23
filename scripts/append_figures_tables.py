"""
Append figures and supplementary tables to manuscript_draft_v1.docx
- Converts each PDF figure to PNG (via PyMuPDF) and embeds it
- Reads each supplementary TSV and adds as a formatted Word table
"""

import fitz  # PyMuPDF
import io
import os
import csv
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCX   = "/Users/nam.tnguyen2022/KIRP_tRF_project/manuscript/manuscript_draft_v1.docx"
FIGS   = "/Users/nam.tnguyen2022/KIRP_tRF_project/results/figures"
TABS   = "/Users/nam.tnguyen2022/KIRP_tRF_project/results/tables"

doc = Document(DOCX)

def set_font(run, bold=False, italic=False, size=12):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic

def add_section_heading(doc, text, level=1):
    h = doc.add_heading('', level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(16 if level == 1 else 13)
    run.bold       = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, italic=True, size=11)

def pdf_to_png_bytes(pdf_path, dpi=150):
    """Render first page of PDF to PNG bytes."""
    pdf = fitz.open(pdf_path)
    page_count = len(pdf)
    images = []
    for page_num in range(page_count):
        page = pdf[page_num]
        mat  = fitz.Matrix(dpi / 72, dpi / 72)
        pix  = page.get_pixmap(matrix=mat, alpha=False)
        images.append(pix.tobytes("png"))
    pdf.close()
    return images

def add_tsv_table(doc, tsv_path, max_rows=50):
    """Read TSV and add as Word table (up to max_rows data rows)."""
    with open(tsv_path, newline='', encoding='utf-8') as f:
        reader = csv.reader(f, delimiter='\t')
        rows = list(reader)
    if not rows:
        return
    header = rows[0]
    data   = rows[1:max_rows+1]
    ncols  = len(header)
    table  = doc.add_table(rows=1 + len(data), cols=ncols)
    table.style = 'Table Grid'
    # Header row
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(h)
        set_font(r, bold=True, size=9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Light blue header shading
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = qn('w:shd')
        from lxml import etree
        shd_elem = etree.SubElement(tcPr, shd)
        shd_elem.set(qn('w:val'),   'clear')
        shd_elem.set(qn('w:color'), 'auto')
        shd_elem.set(qn('w:fill'), 'DCE6F1')
    # Data rows
    for i, row in enumerate(data):
        for j, val in enumerate(row[:ncols]):
            cell = table.cell(i+1, j)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(val)
            set_font(r, size=8)
    if len(rows) - 1 > max_rows:
        note = doc.add_paragraph()
        r = note.add_run(f"  (Showing first {max_rows} of {len(rows)-1} rows. Full table in results/tables/)")
        set_font(r, italic=True, size=9)

# ─────────────────────────────────────────────────────────────────────────────
# FIGURES SECTION
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
add_section_heading(doc, "Figures", level=1)

figures = [
    ("trf_volcano.pdf",
     "Figure 1. Volcano plot of differentially expressed tRFs in KIRP. "
     "Red points: FDR < 0.05 and |log₂FC| ≥ 1. "
     "Labeled: top 10 most significant tRFs."),
    ("trf_heatmap_top50.pdf",
     "Figure 2. Heatmap of the top 50 differentially expressed tRFs. "
     "Columns represent individual samples (tumor vs. normal); rows represent tRFs. "
     "Expression shown as log₂(CPM+1), scaled by row."),
    ("trf_boxplots_top10.pdf",
     "Figure 3. Boxplots of the top 10 differentially expressed tRFs by FDR, "
     "comparing tumor (red) and normal (blue) samples."),
    ("gene_volcano.pdf",
     "Figure 4. Volcano plot of KIRP candidate gene differential expression. "
     "Significant up-regulated genes (MET, CDKN2A, TERT) are labeled."),
    ("gene_heatmap.pdf",
     "Figure 5. Heatmap of KIRP candidate gene expression across tumor and normal samples."),
    ("gene_boxplots_significant.pdf",
     "Figure 6. Boxplots of significantly differentially expressed KIRP candidate genes."),
    ("anticorrelation_heatmap.pdf",
     "Figure 7. Heatmap of Spearman correlation coefficients for the top anticorrelated "
     "tRF–gene pairs (tRF down-regulated, gene up-regulated in tumor)."),
    ("anticorrelation_scatterplots_top.pdf",
     "Figure 8. Scatterplots of top anticorrelated tRF–gene pairs across 290 matched tumor samples. "
     "Each point represents one sample; line shows Spearman fit."),
    ("roc_curves_top5.pdf",
     "Figure 9. ROC curves for the top 5 tRFs distinguishing KIRP tumor from normal tissue. "
     "AUC values are shown in the legend. Dashed diagonal: random classifier."),
    ("roc_auc_barplot.pdf",
     "Figure 10. AUC values for the top 20 differentially expressed tRFs. "
     "Colour gradient indicates AUC magnitude. Dashed line: AUC = 0.5."),
    ("stage_association_boxplots.pdf",
     "Figure 11. tRF expression stratified by primary KIRP AJCC pathologic stage (I–IV) "
     "in 261 tumor samples with stage annotation. "
     "Kruskal–Wallis FDR values are reported in Supplementary Table S7."),
    ("km_survival_top3_trfs.pdf",
     "Figure 12. Kaplan–Meier overall survival curves for patients stratified by "
     "median expression of the top 3 tRFs (by AUC). P-values from log-rank test."),
    ("cox_forest_plot.pdf",
     "Figure 13. Cox univariate hazard ratios (HR, 95% CI) for the top 20 tRFs. "
     "tRFs with Cox FDR < 0.05 are shown in red. Vertical dashed line: HR = 1."),
]

for fig_file, caption in figures:
    fig_path = os.path.join(FIGS, fig_file)
    if not os.path.exists(fig_path):
        print(f"  MISSING: {fig_file}")
        continue
    print(f"  Embedding: {fig_file}")
    pages = pdf_to_png_bytes(fig_path, dpi=150)
    for page_png in pages:
        img_stream = io.BytesIO(page_png)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(6.0))
    add_caption(doc, caption)
    doc.add_paragraph()  # spacing

# ─────────────────────────────────────────────────────────────────────────────
# SUPPLEMENTARY TABLES SECTION
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
add_section_heading(doc, "Supplementary Tables", level=1)

supp_tables = [
    ("trf_DE_significant.tsv",
     "Supplementary Table S1. Differentially expressed tRFs in KIRP "
     "(FDR < 0.05, |log₂FC| ≥ 1; 1,585 tRFs). "
     "First 50 rows shown; full table in results/tables/trf_DE_significant.tsv.",
     50),
    ("gene_DE_significant.tsv",
     "Supplementary Table S2. Differentially expressed KIRP candidate genes "
     "(FDR < 0.05, |log₂FC| ≥ 1).",
     50),
    ("anticorrelation_prioritized.tsv",
     "Supplementary Table S3. Significant anticorrelated tRF–gene pairs "
     "(tRF down-regulated, gene up-regulated in tumor; FDR < 0.05). "
     "First 50 of 104 pairs shown.",
     50),
    ("rna22_binding_summary.tsv",
     "Supplementary Table S4. RNA22 binding predictions — best hit per tRF–gene pair "
     "(p < 0.05; 18 pairs).",
     50),
    ("roc_auc_top_trfs.tsv",
     "Supplementary Table S5. AUC values for the top 20 tRFs distinguishing "
     "KIRP tumor from normal tissue.",
     50),
    ("cox_univariate_results.tsv",
     "Supplementary Table S6. Cox univariate hazard ratios for the top 20 tRFs "
     "and overall survival.",
     50),
    ("stage_kruskal_results.tsv",
     "Supplementary Table S7. Kruskal–Wallis test results for tRF expression "
     "association with AJCC pathologic stage.",
     50),
]

for tsv_file, caption, max_rows in supp_tables:
    tsv_path = os.path.join(TABS, tsv_file)
    if not os.path.exists(tsv_path):
        print(f"  MISSING: {tsv_file}")
        continue
    print(f"  Adding table: {tsv_file}")
    add_caption(doc, caption)
    add_tsv_table(doc, tsv_path, max_rows=max_rows)
    doc.add_paragraph()

doc.save(DOCX)
print(f"\n✓ Updated document saved → {DOCX}")
