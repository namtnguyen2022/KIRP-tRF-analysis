"""
Regenerate manuscript_draft_v1.docx from updated markdown,
then append 6 composite main figures + supplementary tables.
"""

import fitz
import io, os, csv, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

MD    = "/Users/nam.tnguyen2022/KIRP_tRF_project/manuscript/manuscript_draft_v1.md"
OUT   = "/Users/nam.tnguyen2022/KIRP_tRF_project/manuscript/manuscript_draft_v1.docx"
FIGS  = "/Users/nam.tnguyen2022/KIRP_tRF_project/results/figures"
TABS  = "/Users/nam.tnguyen2022/KIRP_tRF_project/results/tables"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=12):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(doc, text, level):
    h = doc.add_heading('', level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt({1:16,2:14,3:12,4:12}.get(level,12))
    run.bold = True
    run.font.color.rgb = RGBColor(0,0,0)

def parse_inline(para, text):
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()]); set_font(r)
        token = m.group()
        if token.startswith('***'):
            r = para.add_run(token[3:-3]); set_font(r, bold=True, italic=True)
        elif token.startswith('**'):
            r = para.add_run(token[2:-2]); set_font(r, bold=True)
        else:
            r = para.add_run(token[1:-1]); set_font(r, italic=True)
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:]); set_font(r)

def add_table_from_md(doc, lines):
    rows = [l for l in lines if l.startswith('|')]
    data_rows = [r for r in rows if not re.match(r'^\|[-| :]+\|$', r)]
    if not data_rows: return
    ncols = data_rows[0].count('|') - 1
    table = doc.add_table(rows=len(data_rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row_text in enumerate(data_rows):
        cells = [c.strip() for c in row_text.strip('|').split('|')]
        for j, cell_text in enumerate(cells[:ncols]):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]; p.clear()
            parts = re.split(r'(\*[^*]+\*)', cell_text)
            for part in parts:
                if part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1]); set_font(r, italic=True, bold=(i==0))
                else:
                    r = p.add_run(part); set_font(r, bold=(i==0), size=10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_main_table(doc, tsv_path, label, caption, cols=None, max_rows=None,
                   col_widths=None, left_align_cols=None):
    """Add a numbered main table with bold label + italic caption, then the table."""
    with open(tsv_path, newline='', encoding='utf-8') as f:
        rows = list(csv.reader(f, delimiter='\t'))
    if not rows:
        return
    header = rows[0]
    if cols:
        idx = [header.index(c) for c in cols if c in header]
        header = [header[i] for i in idx]
        data   = [[r[i] for i in idx] for r in rows[1:]]
    else:
        data = rows[1:]
    if max_rows:
        data = data[:max_rows]

    ncols = len(header)

    # Bold table label + italic caption on same paragraph
    p = doc.add_paragraph()
    r_lbl = p.add_run(label + ". "); set_font(r_lbl, bold=True, size=10)
    r_cap = p.add_run(caption);      set_font(r_cap, italic=True, size=10)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)

    table = doc.add_table(rows=1 + len(data), cols=ncols)
    table.style = 'Table Grid'

    # Set column widths if provided
    if col_widths:
        for j, w in enumerate(col_widths[:ncols]):
            for row in table.rows:
                row.cells[j].width = Inches(w)

    # Header row — blue shading, bold, centred
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        p2 = cell.paragraphs[0]; p2.clear()
        r2 = p2.add_run(h); set_font(r2, bold=True, size=8.5)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = etree.SubElement(tcPr, qn('w:shd'))
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')

    # Data rows — alternating light shading on even rows
    for i, row in enumerate(data):
        fill = 'F2F2F2' if i % 2 == 1 else 'FFFFFF'
        for j, val in enumerate(row[:ncols]):
            cell = table.cell(i + 1, j)
            p2 = cell.paragraphs[0]; p2.clear()
            # italicise gene names (all-caps 2-6 letters common in these tables)
            is_gene_col = j == 0 and re.match(r'^[A-Z][A-Z0-9]{1,5}$', val.strip())
            r2 = p2.add_run(val)
            set_font(r2, italic=bool(is_gene_col), size=8)
            # left-align specified columns (or all non-centre cols)
            if left_align_cols and j in left_align_cols:
                p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = etree.SubElement(tcPr, qn('w:shd'))
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); set_font(r, italic=True, size=10)

def pdf_pages_as_png(pdf_path, dpi=150):
    pdf = fitz.open(pdf_path)
    imgs = []
    for page in pdf:
        pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72), alpha=False)
        imgs.append(pix.tobytes("png"))
    pdf.close()
    return imgs

def add_tsv_table(doc, tsv_path, max_rows=50):
    with open(tsv_path, newline='', encoding='utf-8') as f:
        rows = list(csv.reader(f, delimiter='\t'))
    if not rows: return
    header, data = rows[0], rows[1:max_rows+1]
    ncols = len(header)
    table = doc.add_table(rows=1+len(data), cols=ncols)
    table.style = 'Table Grid'
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(h); set_font(r, bold=True, size=8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = etree.SubElement(tcPr, qn('w:shd'))
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'DCE6F1')
    for i, row in enumerate(data):
        for j, val in enumerate(row[:ncols]):
            cell = table.cell(i+1, j)
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(val); set_font(r, size=7)
    if len(rows)-1 > max_rows:
        note = doc.add_paragraph()
        r = note.add_run(f"  (First {max_rows} of {len(rows)-1} rows shown)")
        set_font(r, italic=True, size=8)

# ── Build main text from markdown ────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(2.5)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

with open(MD, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')
    if re.match(r'^---+$', line.strip()): i += 1; continue
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        add_heading(doc, m.group(2).strip(), len(m.group(1))); i += 1; continue
    if line.startswith('|'):
        tbl = []
        while i < len(lines) and lines[i].startswith('|'):
            tbl.append(lines[i].rstrip('\n')); i += 1
        add_table_from_md(doc, tbl); continue
    m = re.match(r'^[-*]\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Bullet')
        parse_inline(p, m.group(1)); i += 1; continue
    m = re.match(r'^\d+\.\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        parse_inline(p, m.group(1)); i += 1; continue
    if not line.strip(): i += 1; continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parse_inline(p, line); i += 1

# ── MAIN FIGURES ─────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Figures", level=1)

main_figures = [
    (
        ["fig2_trf_DE.pdf"],
        "Figure 1. Differential expression of tRFs in KIRP. "
        "(A) Volcano plot of 5,000 tested tRFs (red = up-regulated, blue = down-regulated; FDR<0.05, |log2FC|>=1). "
        "(B) Heatmap of the top 30 tRFs by |log2FC|, scaled by row. "
        "(C) Boxplots of the top 9 tRFs by FDR in tumor vs. normal tissue."
    ),
    (
        ["fig3_gene_DE.pdf"],
        "Figure 2. Differential expression of KIRP candidate target genes. "
        "(A) Volcano plot of 16 candidate genes ranked by log₂FC. Dot size = −log₁₀(FDR); "
        "* FDR<0.05, ** FDR<0.01, *** FDR<0.001; dashed lines at |log₂FC|=1. "
        "(B–D) Individual sample expression (filled dots) for MET, CDKN2A, and TERT in tumor vs. solid normal tissue. "
        "Box = IQR + whiskers; bracket = Wilcoxon test."
    ),
    (
        ["fig4_trf_met_axis.pdf"],
        "Figure 3. The tRF-MET regulatory axis in KIRP. "
        "(A-E) Spearman correlation of the top 5 anticorrelated tRF-MET pairs across 290 matched tumor samples. "
        "(F) RNA22-predicted binding sites on the MET 3'UTR; dot size = -log10(p), red = p<0.05."
    ),
    (
        ["fig5_roc.pdf"],
        "Figure 4. Internal tumor-normal discrimination by candidate tRFs. "
        "(A) ROC curves for the top 5 tRFs distinguishing KIRP tumor from solid normal tissue (internal cohort). "
        "(B) AUC lollipop for the top 20 differentially expressed tRFs (red = AUC >= 0.90). "
        "External validation is required before any clinical interpretation."
    ),
    (
        ["fig6_stage_survival.pdf"],
        "Figure 5. Clinical associations of tRF expression in KIRP. "
        "(A) All five significant candidate stage-associated tRFs stratified by primary KIRP AJCC pathologic stage. "
        "(B) Kaplan-Meier overall survival by detected versus not detected tRF-30-81PV6RRNLNK8 expression "
        "(univariable Cox HR=1.16, FDR=0.033)."
    ),
]

for fig_files, caption in main_figures:
    for fig_file in fig_files:
        fig_path = os.path.join(FIGS, fig_file)
        if not os.path.exists(fig_path):
            print(f"  MISSING: {fig_file}"); continue
        print(f"  Embedding: {fig_file}")
        for png in pdf_pages_as_png(fig_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(png), width=Inches(5.8))
    add_caption(doc, caption)
    doc.add_paragraph()

# ── MAIN TABLES ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Tables", level=1)

# Table 1 — Cohort and sample characteristics
add_main_table(
    doc,
    tsv_path        = os.path.join(TABS, "table1_cohort.tsv"),
    label           = "Table 1",
    caption         = ("Cohort and sample characteristics for TCGA-KIRP. "
                       "Clinical data derived from GDC Data Portal (accessed May 2026). "
                       "tRF = transfer RNA-derived fragment; mRNA = messenger RNA; "
                       "IQR = interquartile range; AJCC = American Joint Committee on Cancer."),
    max_rows        = None,
    col_widths      = [3.8, 2.8],
    left_align_cols = [0],
)

# Table 2 — Prioritised candidate tRFs by evidence category
add_main_table(
    doc,
    tsv_path        = os.path.join(TABS, "table2_candidates.tsv"),
    label           = "Table 2",
    caption         = ("Prioritised candidate tRF biomarkers and regulatory pairs in TCGA-KIRP. "
                       "Each candidate is presented with its primary evidence category and key supporting statistics. "
                       "AUC = area under ROC curve (tumor vs. solid normal tissue; internal cohort only). "
                       "Stage FDR = Benjamini\u2013Hochberg-adjusted Kruskal\u2013Wallis p-value across AJCC stages I\u2013IV. "
                       "\u1d43 Exploratory univariable association only; no adjustment for age, sex, or stage. "
                       "Not equivalent to independent prognostic evidence."),
    max_rows        = None,
    col_widths      = [1.7, 1.5, 0.7, 2.8, 1.9],
    left_align_cols = [3, 4],
)

# ── SUPPLEMENTARY TABLES — captions only (data uploaded separately) ──────────
doc.add_page_break()
add_heading(doc, "Supplementary Tables", level=1)

supp_captions = [
    ("Supplementary Table S1. "
     "Significantly differentially expressed tRFs in TCGA-KIRP (edgeR; FDR\u202f<\u202f0.05, "
     "|log\u2082FC|\u202f\u2265\u202f1; n\u202f=\u202f1,585). "
     "Full table provided as Supplementary Table S1."),
    ("Supplementary Table S2. "
     "Differential expression of KIRP candidate genes (edgeR; FDR\u202f<\u202f0.05). "
     "logFC\u202f=\u202flog\u2082 fold change (Tumor/Normal)."),
    ("Supplementary Table S3. "
     "Significant anticorrelated tRF\u2013gene pairs in KIRP tumor samples "
     "(Spearman r\u202f<\u202f0, BH-FDR\u202f<\u202f0.05; n\u202f=\u202f290 matched samples; 104 pairs total). "
     "Full table provided as Supplementary Table S3."),
    ("Supplementary Table S4. "
     "Full RNA22 v2 predicted tRF\u20133\u02b9UTR binding interactions including heteroduplex structures "
     "(p\u202f<\u202f0.05; 18 unique tRF\u2013gene pairs). "
     "Full table provided as Supplementary Table S4."),
    ("Supplementary Table S5. "
     "Internal tumor-normal discrimination performance (AUC) for the top 20 differentially expressed tRFs "
     "(pROC; TCGA-KIRP cohort only)."),
    ("Supplementary Table S6. "
     "Exploratory univariable Cox proportional-hazards regression for overall survival "
     "(n\u202f=\u202f251 tumor samples; 44 events). "
     "Each tRF modelled as sole continuous predictor; no covariate adjustment."),
    ("Supplementary Table S7. "
     "Kruskal\u2013Wallis test for tRF expression across primary KIRP AJCC pathologic stages I\u2013IV "
     "(n\u202f=\u202f261 annotated tumor samples; BH-FDR adjusted)."),
]

for caption in supp_captions:
    p = doc.add_paragraph()
    r = p.add_run(caption)
    set_font(r, size=10)
    p.paragraph_format.space_after = Pt(6)

doc.save(OUT)
print(f"\n\u2713 Document saved -> {OUT}")
