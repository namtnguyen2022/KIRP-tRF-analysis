"""
Convert manuscript_draft_v1.md to a formatted Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD   = "/Users/nam.tnguyen2022/KIRP_tRF_project/manuscript/manuscript_draft_v1.md"
OUT  = "/Users/nam.tnguyen2022/KIRP_tRF_project/manuscript/manuscript_draft_v1.docx"

doc = Document()

# ── Page margins (2.5 cm all sides) ─────────────────────────────────────────
from docx.shared import Cm
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default body style ───────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def set_font(run, bold=False, italic=False, size=12, color=None):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level):
    h = doc.add_heading('', level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12, 4: 12}
    run.font.name = 'Times New Roman'
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_table_from_md(doc, lines):
    """Parse a markdown table and add a Word table."""
    rows = [l for l in lines if l.startswith('|')]
    if len(rows) < 2:
        return
    # Remove separator row (---|---)
    data_rows = [r for r in rows if not re.match(r'^\|[-| :]+\|$', r)]
    if not data_rows:
        return
    ncols = data_rows[0].count('|') - 1
    table = doc.add_table(rows=len(data_rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row_text in enumerate(data_rows):
        cells = [c.strip() for c in row_text.strip('|').split('|')]
        for j, cell_text in enumerate(cells[:ncols]):
            cell = table.cell(i, j)
            # Apply italic for gene names (*GENE*)
            p = cell.paragraphs[0]
            p.clear()
            parts = re.split(r'(\*[^*]+\*)', cell_text)
            for part in parts:
                if part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    set_font(r, italic=True, bold=(i == 0))
                else:
                    r = p.add_run(part)
                    set_font(r, bold=(i == 0))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def parse_inline(para, text):
    """Add text to paragraph with bold (**) and italic (*) formatting."""
    # Patterns: **bold**, *italic*, ***bold-italic***, superscripts (²⁶ etc.)
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)')
    pos = 0
    for m in pattern.finditer(text):
        # Plain text before match
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            set_font(r)
        token = m.group()
        if token.startswith('***'):
            r = para.add_run(token[3:-3])
            set_font(r, bold=True, italic=True)
        elif token.startswith('**'):
            r = para.add_run(token[2:-2])
            set_font(r, bold=True)
        else:
            r = para.add_run(token[1:-1])
            set_font(r, italic=True)
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        set_font(r)

# ── Read and parse markdown ──────────────────────────────────────────────────
with open(MD, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip horizontal rules
    if re.match(r'^---+$', line.strip()):
        i += 1
        continue

    # Headings
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = m.group(2).strip()
        add_heading(doc, text, level)
        i += 1
        continue

    # Table: collect all consecutive table lines
    if line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].startswith('|'):
            table_lines.append(lines[i].rstrip('\n'))
            i += 1
        add_table_from_md(doc, table_lines)
        continue

    # Bullet list
    m = re.match(r'^[-*]\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Bullet')
        parse_inline(p, m.group(1))
        for run in p.runs:
            set_font(run)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^\d+\.\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        parse_inline(p, m.group(1))
        i += 1
        continue

    # Bold italic title line (manuscript title at top)
    if line.startswith('# '):
        add_heading(doc, line[2:], 1)
        i += 1
        continue

    # Author / affiliation lines (lines starting with **)
    if line.startswith('**') and line.endswith('**') and len(line) < 100:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.strip('*'))
        set_font(r, bold=True)
        i += 1
        continue

    # Superscript affiliation lines (¹ etc.)
    if line and line[0] in '¹²³⁴⁵':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_font(r, italic=True, size=10)
        i += 1
        continue

    # Empty line → paragraph break (skip, docx handles spacing)
    if not line.strip():
        i += 1
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parse_inline(p, line)
    i += 1

doc.save(OUT)
print(f"Word document saved → {OUT}")
